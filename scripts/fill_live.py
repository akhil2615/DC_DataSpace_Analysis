"""Fill the Data Cloud Data Space Analysis Record from the live cache.

Run scripts/fetch_live.py first. Nothing in here is hand-typed org data: every
cell traces to a cached endpoint response.

Conventions
  * A cell is left BLANK when the column is a migration decision or a human input.
    Every blank column is listed in the "Human input register" appendix with where to
    get the answer and whether it needs a second reviewer.
  * A cell reads NOT AVAILABLE FROM API when the column is factual but no endpoint
    exposes it (full-refresh duration, partitioning, IR run duration, ...).
  * Derived cells state what they were derived from, in the provenance appendix.
  * The template file is never modified; output goes to a new file.
"""

from __future__ import annotations

import html
import json
import re
from datetime import date
from pathlib import Path

import docx

REPO = Path(__file__).resolve().parent.parent
CACHE = REPO / ".data-space-analysis-cache"
TEMPLATE = REPO / "DataCloud_DataSpace_Analysis_Record_TEMPLATE.docx"

NA = "NOT AVAILABLE FROM API"
NONE_DETECTED = "none detected"

CONNECTOR_UI_LABELS = {
    "SalesforceDotCom": "Salesforce CRM",
    "SalesforceCRM": "Salesforce CRM",
    "SalesforceMarketingCloud": "Salesforce Marketing Cloud",
    "IngestApi": "Ingestion API",
    "StreamingApp": "Website and Mobile App",
    "UploadedFiles": "File Upload",
    "SFTP": "Secure File Transfer (SFTP)",
    "BIGQUERY": "Google Big Query",
    "AwsS3": "Amazon S3",
    "FacebookAds": "Meta Ads",
    "MarketingCloudPersonalization": "Marketing Cloud Personalization",
    "SalesforceCommerceCloud": "Commerce Cloud",
}

CARDINALITY = {
    "NTOONE": "N to One",
    "ONETON": "One to N",
    "ONETOONE": "One to One",
    "NTON": "N to N",
}

# Formula targets every stream carries because the platform adds them.
SYSTEM_FORMULA_TARGETS = {"DataSource", "DataSourceObject", "cdp_sys_PartitionDate"}

# Labels that look provisional. Flagged for a human to confirm, never acted on.
PROVISIONAL_LABEL = re.compile(r"^\s*\d+\s*$|test|temp|demo|copy|dummy|asdf", re.I)


# ---------------------------------------------------------------- cache access


def load(name: str) -> dict:
    p = CACHE / f"{name}.json"
    if not p.exists():
        return {}
    return json.loads(p.read_text(encoding="utf-8"))


def records(name: str) -> list:
    r = load(name).get("records")
    return r if isinstance(r, list) else []


# ---------------------------------------------------------------- docx helpers


def ensure_rows(table, needed: int) -> None:
    while len(table.rows) < needed:
        table.add_row()


def put(table, r: int, c: int, v) -> None:
    if r < len(table.rows) and c < len(table.rows[r].cells):
        table.rows[r].cells[c].text = "" if v is None else str(v)


def drop_rows_from(table, start_row: int) -> None:
    """Remove the template's spare placeholder rows once real rows are written."""
    for row in list(table.rows)[start_row:]:
        table._tbl.remove(row._tr)


def write_rows(table, rows: list[tuple], start: int = 1) -> None:
    ensure_rows(table, start + len(rows))
    for i, row in enumerate(rows, start=start):
        for c, v in enumerate(row):
            put(table, i, c, v)
    drop_rows_from(table, start + len(rows))


def add_table(doc, headers: list[str], rows: list[tuple]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for c, h in enumerate(headers):
        t.rows[0].cells[c].text = h
    for row in rows:
        cells = t.add_row().cells
        for c, v in enumerate(row):
            if c < len(cells):
                cells[c].text = "" if v is None else str(v)
    return t


# ---------------------------------------------------------------- value shaping


def ui_label(raw: str) -> str:
    return CONNECTOR_UI_LABELS.get(raw, raw or "")


def cardinality(raw: str) -> str:
    return CARDINALITY.get((raw or "").upper(), raw or "")


def object_type(api_name: str) -> str:
    if api_name.endswith("__dll"):
        return "DLO"
    if api_name.endswith("__dlm"):
        return "DMO"
    if api_name.endswith("__cio"):
        return "Calculated insight"
    return NA


def unescape_json(raw: str):
    """Segment criteria arrive HTML-escaped; return parsed JSON or None."""
    if not raw:
        return None
    try:
        return json.loads(html.unescape(raw))
    except (ValueError, TypeError):
        return None


def describe_criteria(node, depth: int = 0) -> str:
    """Flatten a segment criteria tree into readable text, no interpretation."""
    if node is None:
        return ""
    if isinstance(node, list):
        return " ; ".join(x for x in (describe_criteria(n, depth) for n in node) if x)
    if not isinstance(node, dict):
        return str(node)

    for group_key in ("operands", "criteria", "nodes", "children", "conditions"):
        if isinstance(node.get(group_key), list):
            joiner = node.get("conjunction") or node.get("operator") or "AND"
            parts = [describe_criteria(n, depth + 1) for n in node[group_key]]
            parts = [p for p in parts if p]
            inner = f" {joiner} ".join(parts)
            return f"({inner})" if depth and len(parts) > 1 else inner

    subject = node.get("subject") or {}
    obj = subject.get("objectApiName") or node.get("objectApiName") or ""
    fld = subject.get("fieldApiName") or node.get("fieldApiName") or ""
    op = node.get("operator") or node.get("type") or ""
    val = node.get("value")
    if val is None:
        val = node.get("values")
    if isinstance(val, (list, dict)):
        val = json.dumps(val, separators=(",", ":"))
    left = ".".join(x for x in (obj, fld) if x)
    text = " ".join(str(x) for x in (left, op, val) if x not in (None, ""))
    return text.strip()


def criteria_objects(node, found: set[str] | None = None) -> set[str]:
    """Every objectApiName referenced anywhere in a criteria tree."""
    found = set() if found is None else found
    if isinstance(node, dict):
        for k, v in node.items():
            if k == "objectApiName" and isinstance(v, str):
                found.add(v)
            else:
                criteria_objects(v, found)
    elif isinstance(node, list):
        for v in node:
            criteria_objects(v, found)
    return found


def schema_leaves(raw: str) -> list[str]:
    """Attribute list from an activation record schema, as object.field pairs."""
    parsed = unescape_json(raw)
    out: list[str] = []

    def walk(node):
        if isinstance(node, dict):
            if node.get("fieldApiName"):
                out.append(f"{node.get('objectApiName','')}.{node['fieldApiName']}".strip("."))
            for k, v in node.items():
                if k not in ("fieldApiName", "objectApiName"):
                    walk(v)
        elif isinstance(node, list):
            for v in node:
                walk(v)

    walk(parsed)
    return sorted(dict.fromkeys(out))


def frequency_text(cfg: dict) -> str:
    """Refresh cadence from refreshConfig, without dumping raw JSON."""
    if not isinstance(cfg, dict) or not cfg:
        return NA
    freq = cfg.get("frequency") if isinstance(cfg.get("frequency"), dict) else {}
    bits = []
    ftype = freq.get("frequencyType")
    if ftype:
        bits.append(str(ftype))
    for key, label in (
        ("hours", "hours"),
        ("refreshDayOfMonth", "day of month"),
        ("daysOfWeek", "days of week"),
        ("minutes", "minutes"),
    ):
        val = freq.get(key)
        if val:
            bits.append(f"{label}={','.join(str(v) for v in val)}")
    mode = cfg.get("refreshMode")
    if mode:
        bits.append(f"mode={mode}")
    return " · ".join(bits) if bits else NA


def walk_graph(node: dict, names: list[str], edges: list[tuple[str, str]]) -> None:
    me = node.get("developerName", "")
    if me:
        names.append(me)
    for child in node.get("relatedObjects") or []:
        kid = child.get("developerName", "")
        if me and kid:
            edges.append((me, kid))
        walk_graph(child, names, edges)


# ---------------------------------------------------------------- human inputs

# section, column, why the API cannot answer it, where the answer comes from, 4-eyes
HUMAN_INPUTS = [
    ("0.2", "Filter must be reproduced?", "Migration decision", "Workshop with data space owner", "Yes"),
    ("0.3", "Evidence file sign-off", "Process control", "Whoever ran the extract", "No"),
    ("1.2", "All questions", "Migration disposition", "Programme decision log D1-D17", "Yes"),
    ("1.3", "Migrate / Drop / DataKit / CLI JSON / Manual UI", "Migration decision", "Disposition workshop", "Yes"),
    ("1.4", "Top three risks", "Judgement", "Architect assessment", "Yes"),
    ("2.1", "Full-refresh duration (h)", "Not exposed by the API", "Time a full refresh in the source org, or read Data Cloud job history in the UI", "No"),
    ("2.1", "Migrate? / Rebuild path", "Migration decision", "Disposition workshop", "Yes"),
    ("2.2", "Owning team", "Not held in Data Cloud", "Connection owner from the CMDB or platform team", "No"),
    ("2.2", "Connection exists in target org? / New setup required?", "Target-org fact", "Run the same extract against the target org", "Yes"),
    ("2.3", "Historical data and backfill", "Source-system fact", "Source system owner plus decision D12", "Yes"),
    ("3.1", "Partitioning", "Not exposed by the API", "Data Cloud UI, DLO detail page", "No"),
    ("3.2", "Data-quality observations", "Needs profiling", "Stage-2 profiling workbook", "Yes"),
    ("4.1", "Subject area, where blank", "Only standard CIM DMOs carry RefEntitySubjectArea; custom and platform-generated ones have none", "Assign a business subject area during the harmonisation workshop", "No"),
    ("4.3", "Field-level disposition counts", "Needs the mapping workbook", "Field-level mapping workbook, stage 4", "Yes"),
    ("4.4", "Attribute precedence", "Business rule", "Data owner workshop; cross-check against 5.1b", "Yes"),
    ("4.5", "Transformations required", "Design decision", "Mapping workbook and transform design", "Yes"),
    ("5.1a", "IR run duration (h)", "Not exposed by the API", "Data Cloud UI, identity resolution job history", "No"),
    ("5.1a", "Boolean logic / Migrate as-is?", "Judgement", "IR owner review", "Yes"),
    ("5.2", "Ruleset target changes", "Design decision", "IR redesign workshop", "Yes"),
    ("5.3", "Template versus country variant", "Programme decision", "Template governance board", "Yes"),
    ("6.1", "Run duration (h)", "Not exposed by the API", "Data Cloud UI, CI run history", "No"),
    ("6.2", "Consumers", "Not exposed by the API", "Ask the Personalization or agent owners", "No"),
    ("7.1", "Business purpose", "Not held in Data Cloud", "Segment owner; the API carries no description field", "No"),
    ("7.1", "Owner", "Not held in Data Cloud", "Marketing owner list; segments carry no createdBy", "No"),
    ("7.2", "Depends on EXT-2/3 field? / Logic changes?", "Depends on 4.3", "Cross-check once 4.3 is filled", "Yes"),
    ("7.3", "Segments to migrate / drop / redefine", "Migration decision", "Segment rationalisation workshop", "Yes"),
    ("8.1a", "Mechanism / CUSTOM sub-type", "Judgement on target setup", "Activation owner plus destination platform team", "Yes"),
    ("8.1b", "of which PII / special category", "Data classification", "Privacy team review of the attribute list in this row", "Yes"),
    ("8.1b", "Consent basis", "Legal basis", "Privacy and legal", "Yes"),
    ("8.1b", "Re-create from Home Org?", "Migration decision", "Activation owner", "Yes"),
    ("8.2", "Destination detail", "External account facts", "Destination platform owners", "Yes"),
    ("8.3", "MC journey dependencies", "Marketing Cloud fact", "MC platform team", "Yes"),
    ("9.1", "Consent model", "Policy", "Privacy team", "Yes"),
    ("9.2", "Residency constraints", "Policy", "Legal and privacy", "Yes"),
    ("10.1", "Consumer inventory", "Partly outside Data Cloud", "Ask platform owners; data actions and data action targets are extracted automatically", "Yes"),
    ("10.2", "Real-time personalisation", "Outside this API", "Personalization app owner", "No"),
    ("11.1", "Direct-read consumers", "Outside Data Cloud", "Integration inventory", "Yes"),
    ("11.2", "Query and SQL consumers", "Outside Data Cloud", "Integration inventory and query logs", "Yes"),
    ("11.3", "Consumer / Re-index duration (h)", "Not exposed by the API", "Agent owner; time a re-index", "No"),
    ("11.4", "Data kit packaging", "ssot/data-kits returns 500 on this org", "Data Cloud UI, Developer Tools, Data Kits", "Yes"),
    ("12.1", "Tolerance / Test method / Signed by", "Test design", "Test lead; baselines are pre-filled from the extract", "Yes"),
    ("12.2", "Divergence triggers", "Judgement", "Architect assessment", "Yes"),
    ("12.3", "Open questions", "Programme tracking", "Architect and decision log", "No"),
    ("12.4", "Build order estimates", "Planning", "Delivery lead", "Yes"),
    ("12.5", "Sign-off", "Governance", "Named approvers", "Yes"),
    ("12.6", "Workbooks and workshop notes", "Produced by humans", "Stages 2 to 4 outputs", "No"),
]


# ---------------------------------------------------------------- section fills


def fill(doc, prov: dict, target_space: str) -> tuple[dict, list[tuple]]:
    t = doc.tables
    for table in t:
        table.style = "Table Grid"

    counts: dict[str, str] = {}
    today = date.today().strftime("%d %b %y")
    fetched = (prov.get("fetchedAtUtc") or "")[:10]
    extract_date = fetched or today

    spaces = records("data-spaces")
    by_name = {s.get("name"): s for s in spaces if s.get("name")}
    if target_space not in by_name:
        known = ", ".join(sorted(by_name)) or "(none in cache)"
        raise ValueError(f"Data space '{target_space}' not found in cache. Available: {known}")
    space = by_name[target_space]
    space_name = target_space

    streams = records("data-streams")
    dlos = records("data-lake-objects")
    catalogue = {d.get("name"): d for d in records("data-model-objects-catalogue")}
    meta_dmo = records("metadata-dmo")
    meta_dlo = {d.get("name"): d for d in records("metadata-dlo")}
    meta_ci = records("metadata-ci")
    ci_items = {c.get("apiName"): c for c in records("calculated-insights")}
    idres = records("identity-resolutions")
    segments = records("segments")
    activations = records("activations")
    targets = {a.get("id"): a for a in records("activation-targets")}
    transforms = records("data-transforms")
    indexes = records("search-index")
    graphs = records("data-graphs")
    conns_by_type = load("connections").get("byType", {})
    mappings = load("dmo-mappings").get("byDmo", {})
    graph_detail = load("data-graph-details").get("byGraph", {})
    act_detail = load("activation-details").get("byActivation", {})
    users = load("users").get("byId", {})
    members = load("data-space-members").get("bySpace", {}).get(space_name, {}).get("members", [])

    def rec_space(rec: dict) -> str | None:
        for k in ("dataSpaceName", "dataSpace", "dataspaceName", "dataspace"):
            v = rec.get(k)
            if isinstance(v, str) and v:
                return v
        return None

    def in_space(rec: dict) -> bool:
        s = rec_space(rec)
        return s is None or s == space_name

    # Enforce data-space scoped records where payload includes a space key.
    streams = [s for s in streams if in_space(s)]
    segments = [s for s in segments if in_space(s)]
    activations = [a for a in activations if in_space(a)]
    idres = [r for r in idres if in_space(r)]
    transforms = [t for t in transforms if in_space(t)]
    indexes = [i for i in indexes if in_space(i)]
    graphs = [g for g in graphs if in_space(g)]
    ci_items = {k: v for k, v in ci_items.items() if in_space(v)}
    targets = {k: v for k, v in targets.items() if in_space(v)}

    # Most inventory endpoints are org-wide and do not include a space key.
    # Scope those by data-space members (DLOs) and then derive DMO scope from mappings.
    member_dlos = {
        m.get("memberName")
        for m in members
        if isinstance(m, dict) and m.get("memberName")
    }
    if member_dlos:
        dlos = [d for d in dlos if d.get("name") in member_dlos]
        meta_dlo = {k: v for k, v in meta_dlo.items() if k in member_dlos}
        streams = [
            s
            for s in streams
            if ((s.get("dataLakeObjectInfo") or {}).get("name") in member_dlos)
        ]

        scoped_dmos: set[str] = set()
        for by_dmo_name, payload in mappings.items():
            body = (payload or {}).get("body") or {}
            for m in body.get("objectSourceTargetMaps") or []:
                src = m.get("sourceEntityDeveloperName")
                dst = m.get("targetEntityDeveloperName") or by_dmo_name
                if src in member_dlos and dst:
                    scoped_dmos.add(dst)
                    scoped_dmos.add(by_dmo_name)
        if scoped_dmos:
            meta_dmo = [d for d in meta_dmo if d.get("name") in scoped_dmos]
            catalogue = {k: v for k, v in catalogue.items() if k in scoped_dmos}
            mappings = {k: v for k, v in mappings.items() if k in scoped_dmos}

    # Keep detailed payloads aligned to already scoped top-level lists.
    scoped_activation_keys = {a.get("id") or a.get("developerName") for a in activations}
    act_detail = {k: v for k, v in act_detail.items() if k in scoped_activation_keys}
    scoped_graph_keys = {g.get("developerName") for g in graphs if g.get("developerName")}
    graph_detail = {k: v for k, v in graph_detail.items() if k in scoped_graph_keys}

    # Calculated-insight metadata is org-wide; keep only items present in scoped CI list.
    scoped_ci_names = set(ci_items) | {v.get("apiName") for v in ci_items.values() if isinstance(v, dict)}
    if scoped_ci_names:
        meta_ci = [c for c in meta_ci if c.get("name") in scoped_ci_names or c.get("apiName") in scoped_ci_names]

    # Connections endpoint is org-wide. Keep connections referenced by scoped streams.
    scoped_conn_types = set()
    scoped_conn_names = set()
    for s in streams:
        ci = s.get("connectorInfo") or {}
        if ci.get("connectorType"):
            scoped_conn_types.add(ci.get("connectorType"))
        details = ci.get("connectorDetails") or {}
        if details.get("name"):
            scoped_conn_names.add(details.get("name"))
    if scoped_conn_types or scoped_conn_names:
        filtered_by_type = {}
        for ctype, rows in conns_by_type.items():
            if scoped_conn_types and ctype not in scoped_conn_types:
                continue
            keep = [
                r
                for r in rows
                if not scoped_conn_names
                or r.get("name") in scoped_conn_names
                or r.get("label") in scoped_conn_names
            ]
            if keep:
                filtered_by_type[ctype] = keep
        conns_by_type = filtered_by_type

    # Subject area and creation type per DMO, from the Tooling catalogue.
    subject_area: dict[str, str] = {}
    tooling_type: dict[str, str] = {}
    for r in records("tooling-dmo"):
        dev = r.get("DeveloperName")
        if not dev:
            continue
        for key in (f"ssot__{dev}__dlm", f"{dev}__dlm"):
            if r.get("RefEntitySubjectArea"):
                subject_area[key] = r["RefEntitySubjectArea"]
            if r.get("CreationType"):
                tooling_type[key] = r["CreationType"]

    observations: list[tuple] = []

    def flag(area, obj, observation, why, action):
        observations.append((area, obj, observation, why, action))

    # ---------- cover block
    put(t[0], 1, 1, f"DS-{space_name}")
    put(t[0], 2, 1, space.get("label", space_name))
    put(t[0], 8, 1, today)
    if len(doc.paragraphs) > 1:
        doc.paragraphs[1].text = f"DS-{space_name}  ·  {space.get('label', space_name)}"

    # ---------- 0.1 Data space object (t4)
    put(t[4], 1, 1, space.get("label", ""))
    put(t[4], 2, 1, space.get("name", ""))
    put(t[4], 3, 1, "default" if space.get("name") == "default" else "custom")
    for row in range(4, 9):
        put(t[4], row, 1, "")
    put(t[4], 1, 2, f"data space id {space.get('id','')} · status {space.get('status','')}")
    counts["0.1 data space"] = "1 data space (ssot/data-spaces)"

    # ---------- 0.2 Data space scope (t5)
    dlo_by_name = {d.get("name"): d for d in dlos}
    scope_rows = []
    filtered_members = []
    for m in members:
        api = m.get("memberName", "")
        d = dlo_by_name.get(api, {})
        filters = []
        joiner = " AND "
        for entry in d.get("dataSpaceInfo", []) or []:
            if not isinstance(entry, dict) or entry.get("name") != space_name:
                continue
            filt = entry.get("filter") or {}
            if filt.get("conjunctiveOperator") == "OrOperator":
                joiner = " OR "
            for c in filt.get("conditions", []) or []:
                filters.append(
                    f"{c.get('tableName','')}.{c.get('fieldName','')} "
                    f"{c.get('operator','')} {c.get('filterValue','')}".strip()
                )
        if filters:
            filtered_members.append((api, joiner.join(filters)))
        scope_rows.append((d.get("label", api), api, object_type(api), joiner.join(filters), ""))
    write_rows(t[5], scope_rows)
    counts["0.2 data space scope"] = (
        f"{len(scope_rows)} members (ssot/data-spaces/{space_name}/members joined to "
        f"ssot/data-lake-objects for the filter); {len(filtered_members)} carry a filter"
    )
    for api, filt in filtered_members:
        flag(
            "0.2 scope",
            api,
            f"Data space filter is applied: {filt}",
            "A filter that is not reproduced in the target changes every downstream row count",
            "Confirm the filter is intentional and reproduce it in the target org",
        )

    # ---------- 2.1 Stream inventory (t10)
    stream_rows = []
    for s in sorted(streams, key=lambda x: x.get("label", "")):
        dlo_info = s.get("dataLakeObjectInfo") or {}
        conn = s.get("connectorInfo") or {}
        details = conn.get("connectorDetails") or {}
        stream_rows.append(
            (
                s.get("label", ""),
                s.get("name", ""),
                details.get("name") or s.get("dataSource") or NA,
                conn.get("connectorType", NA),
                frequency_text(s.get("refreshConfig") or {}),
                NA,
                s.get("totalRecords", NA),
                dlo_info.get("recordModifiedFieldName", NA),
                f"{s.get('lastRefreshDate', NA)} | {s.get('lastRunStatus', NA)}",
                "",
                "",
            )
        )
        label = s.get("label") or s.get("name", "")
        if s.get("lastRunStatus") not in (None, "SUCCESS"):
            flag(
                "2.1 streams",
                label,
                f"Last run status is {s.get('lastRunStatus')}",
                "A stream that is not succeeding may be carrying stale or partial data",
                "Fix or consciously exclude before using its row count as a baseline",
            )
        if not s.get("totalRecords"):
            flag(
                "2.1 streams",
                label,
                "Row volume is zero or absent",
                "An empty stream is often abandoned, but may also be newly built",
                "Confirm whether it is in scope",
            )
        if not dlo_info.get("recordModifiedFieldName"):
            flag(
                "2.1 streams",
                label,
                "No incremental key, so the stream can only full-refresh",
                "Full refresh drives cutover duration and source-system load",
                "Check whether an incremental key can be added in the target org",
            )
        if s.get("dataAccessMode") == "DIRECT_ACCESS":
            flag(
                "2.1 streams",
                label,
                "Zero-copy stream (dataAccessMode DIRECT_ACCESS)",
                "Zero-copy objects are not migrated like ingested ones; the external share must exist first",
                "Plan the external grant in the target org before deploying dependent objects",
            )
        if PROVISIONAL_LABEL.search(label):
            flag(
                "2.1 streams",
                label,
                "Name looks provisional",
                "Provisional names usually mean experiments that should not migrate",
                "Confirm with the stream owner",
            )
    write_rows(t[10], stream_rows)
    counts["2.1 stream inventory"] = (
        f"{len(stream_rows)} streams (ssot/data-streams?limit=200&includeMappings=true, one call)"
    )

    # ---------- 2.2 Source connections (t11)
    conn_rows = []
    flat = [(ty, c) for ty, rows in conns_by_type.items() for c in rows]
    for ty, c in sorted(flat, key=lambda x: (ui_label(x[0]), x[1].get("label", ""))):
        note = [f"API connectorType: {ty}"]
        bus = [
            b.get("label")
            for b in (c.get("activationBusinessUnits") or [])
            if isinstance(b, dict) and b.get("label")
        ]
        if bus:
            note.append("Activation business units: " + "; ".join(bus))
        spaces_bu = [
            f"{b.get('mid','')}->{b.get('dataspaceName','')}"
            for b in (c.get("businessUnitsToDataSpaces") or [])
            if isinstance(b, dict)
        ]
        if spaces_bu:
            note.append("BU to data space: " + "; ".join(spaces_bu))
        conn_rows.append(
            (c.get("label", ""), c.get("name", ""), ui_label(ty), "", "", "", " | ".join(note))
        )
    write_rows(t[11], conn_rows)
    counts["2.2 source connections"] = (
        f"{len(conn_rows)} connections across {len(conns_by_type)} connector types "
        f"(ssot/connections?connectorType=..., one call per probed type)"
    )

    # ---------- 3.1 DLO inventory (t13)
    records_by_dlo: dict[str, int] = {}
    streams_by_dlo: dict[str, list[str]] = {}
    zero_copy_dlos: set[str] = set()
    for s in streams:
        nm = (s.get("dataLakeObjectInfo") or {}).get("name")
        tot = s.get("totalRecords")
        if nm and isinstance(tot, (int, float)):
            records_by_dlo[nm] = records_by_dlo.get(nm, 0) + int(tot)
        if nm and s.get("name"):
            streams_by_dlo.setdefault(nm, []).append(s["name"])
        if nm and s.get("dataAccessMode") == "DIRECT_ACCESS":
            zero_copy_dlos.add(nm)

    # Formula fields ride along on the stream list response now.
    formulas_by_dlo: dict[str, list[str]] = {}
    for s in streams:
        nm = (s.get("dataLakeObjectInfo") or {}).get("name")
        if not nm:
            continue
        found = formulas_by_dlo.setdefault(nm, [])
        for m in s.get("mappings") or []:
            tgt = m.get("targetFieldName")
            if "formula" in m and tgt and tgt not in SYSTEM_FORMULA_TARGETS:
                found.append(tgt)

    dlo_rows = []
    total_dlo_records = 0
    for d in sorted(dlos, key=lambda x: x.get("label", "")):
        name = d.get("name", "")
        fields = d.get("fields") or d.get("dataLakeFieldInfoRepresentation") or []
        pk = [f.get("name") for f in fields if f.get("isPrimaryKey")]
        if not pk:
            pk = [p.get("name") for p in (meta_dlo.get(name, {}).get("primaryKeys") or [])]
        if name in streams_by_dlo:
            fx = sorted(dict.fromkeys(formulas_by_dlo.get(name, [])))
            formula_cell = f"Y ({len(fx)}): " + "; ".join(fx) if fx else "N"
        else:
            formula_cell = NA
        rec = records_by_dlo.get(name, NA)
        if isinstance(rec, int):
            total_dlo_records += rec
        dlo_rows.append(
            (
                d.get("label", name),
                name,
                d.get("category", NA),
                len(fields) if fields else NA,
                rec,
                ", ".join(x for x in pk if x) or NA,
                formula_cell,
                NA,
                "",
                "",
            )
        )
        if not pk:
            flag(
                "3.1 DLOs",
                name,
                "No primary key declared",
                "Without a primary key, identity resolution and incremental loads behave differently",
                "Confirm whether this is intentional",
            )
        if name not in streams_by_dlo:
            flag(
                "3.1 DLOs",
                name,
                "No data stream feeds this DLO",
                "Usually a platform-generated DLO, but can be an orphan left by a deleted stream",
                "Confirm it is platform-generated, otherwise drop it",
            )
    write_rows(t[13], dlo_rows)
    counts["3.1 dlo inventory"] = (
        f"{len(dlo_rows)} DLOs (ssot/data-lake-objects); "
        f"{sum(1 for r in dlo_rows if r[4] != NA)} row counts joined from the stream that feeds "
        f"each DLO on dataLakeObjectInfo.name; formula fields from the stream mappings, excluding "
        f"the platform-generated targets DataSource, DataSourceObject and cdp_sys_PartitionDate"
    )

    # ---------- 4.1 DMO inventory (t15)
    source_dlos: dict[str, list[str]] = {}
    mapped_fields: dict[str, set[str]] = {}
    for dmo_name, res in mappings.items():
        if res.get("status") != 200:
            continue
        maps = (res.get("body") or {}).get("objectSourceTargetMaps") or []
        srcs = []
        for m in maps:
            for key in (
                "sourceEntityDeveloperName",
                "sourceObjectDeveloperName",
                "sourceEntityName",
                "sourceObjectName",
                "sourceDeveloperName",
            ):
                if m.get(key):
                    srcs.append(m[key])
                    break
            else:
                hit = re.match(r"^(.*?)_map_", m.get("developerName", ""))
                if hit:
                    srcs.append(hit.group(1) + "__dll")
            for f in m.get("fieldMappings") or []:
                if f.get("targetFieldDeveloperName"):
                    mapped_fields.setdefault(dmo_name, set()).add(f["targetFieldDeveloperName"])
        if srcs:
            source_dlos[dmo_name] = sorted(dict.fromkeys(srcs))

    zero_copy_dmos = {
        dmo
        for dmo, srcs in source_dlos.items()
        if any(s in zero_copy_dlos for s in srcs)
    }

    dmo_rows = []
    for d in sorted(meta_dmo, key=lambda x: x.get("displayName", "")):
        name = d.get("name", "")
        cat_fields = (catalogue.get(name) or {}).get("fields") or []
        total = len(cat_fields) if cat_fields else None
        mapped = len(mapped_fields.get(name, ()))
        dmo_rows.append(
            (
                d.get("displayName", name),
                name,
                (catalogue.get(name) or {}).get("creationType") or tooling_type.get(name, NA),
                subject_area.get(name, ""),
                "; ".join(source_dlos.get(name, [])),
                mapped,
                (total - mapped) if total is not None else NA,
                "",
                "",
            )
        )
        if mapped == 0:
            flag(
                "4.1 DMOs",
                name,
                "Deployed but no DLO mapping",
                "Either platform-generated (identity, activation audience, segment membership) or an unused model object",
                "Confirm it is platform-generated, otherwise exclude it from the migration",
            )
    write_rows(t[15], dmo_rows)
    counts["4.1 dmo inventory"] = (
        f"{len(dmo_rows)} deployed DMOs (ssot/metadata?entityType=DataModelObject); "
        f"{len(source_dlos)} have a DLO mapping (ssot/data-model-object-mappings); mapped field "
        f"count is the distinct mapped target fields, total field count comes from "
        f"ssot/data-model-objects, and Subject area is RefEntitySubjectArea from the Tooling "
        f"object MktDataModelObject"
    )

    # ---------- 4.2 DMO relationships (t16)
    graph_edges: dict[tuple[str, str], list[str]] = {}
    graph_objects: dict[str, set[str]] = {}
    for g in graphs:
        names: list[str] = []
        edges: list[tuple[str, str]] = []
        walk_graph(g.get("dgObject") or {}, names, edges)
        graph_objects[g.get("developerName", "")] = set(names)
        for e in edges:
            graph_edges.setdefault(e, []).append(g.get("developerName", ""))
            graph_edges.setdefault(e[::-1], []).append(g.get("developerName", ""))

    seg_parsed = []
    for s in segments:
        inc = unescape_json(s.get("includeCriteria"))
        exc = unescape_json(s.get("excludeCriteria"))
        seg_parsed.append(
            {
                "record": s,
                "include": inc,
                "exclude": exc,
                "objects": criteria_objects(inc) | criteria_objects(exc),
            }
        )

    seen = set()
    rel_rows = []
    used_rel = 0
    for m in meta_dmo:
        for r in m.get("relationships") or []:
            key = (
                r.get("fromEntity"),
                r.get("fromEntityAttribute"),
                r.get("toEntity"),
                r.get("toEntityAttribute"),
                r.get("cardinality"),
            )
            if key in seen:
                continue
            seen.add(key)
            pair = (r.get("fromEntity"), r.get("toEntity"))
            used = [f"graph: {g}" for g in dict.fromkeys(graph_edges.get(pair, []))]
            for p in seg_parsed:
                if pair[0] in p["objects"] and pair[1] in p["objects"]:
                    used.append(f"segment: {p['record'].get('apiName','')}")
            if used:
                used_rel += 1
            rel_rows.append(
                (
                    f"{r.get('fromEntity','')}.{r.get('fromEntityAttribute','')}",
                    f"{r.get('toEntity','')}.{r.get('toEntityAttribute','')}",
                    cardinality(r.get("cardinality", "")),
                    "; ".join(used) or NONE_DETECTED,
                    "",
                    "",
                )
            )
    rel_rows.sort(key=lambda x: (x[0], x[1]))
    write_rows(t[16], rel_rows)
    counts["4.2 dmo relationships"] = (
        f"{len(rel_rows)} unique relationships (ssot/metadata?entityType=DataModelObject); "
        f"'Used by' is derived by matching the from-entity and to-entity pair against the data "
        f"graph object tree and against the objects referenced in segment criteria, so "
        f"{used_rel} are shown as used and the rest read '{NONE_DETECTED}', which means no graph "
        f"or segment we can see references them, not that they are unused"
    )

    # ---------- 5.1a ruleset header (t20)
    header_rows = []
    for rs in idres:
        header_rows.extend(
            [
                ("Ruleset name / API name", f"{rs.get('label','')} / {rs.get('id','')}"),
                ("Applies to DMO", rs.get("secondaryDmo", rs.get("objectApiName", ""))),
                ("Individuals in", rs.get("sourceProfiles", "")),
                ("Unified profiles out", rs.get("totalUnifiedProfiles", "")),
                (
                    "Match rate (measured)",
                    f"{rs.get('consolidationRate','')} % consolidation "
                    f"(as at {rs.get('lastJobCompleted','')})",
                ),
                (
                    "Last successful IR run",
                    f"{rs.get('lastJobCompleted','')} | {rs.get('lastJobStatus','')}",
                ),
                ("IR run duration", NA),
                ("Ruleset status", rs.get("rulesetStatus", "")),
                ("Runs automatically", rs.get("doesRunAutomatically", "")),
                ("Matched source profiles", rs.get("matchedSourceProfiles", "")),
                ("Known unified profiles", rs.get("knownUnifiedProfiles", "")),
                ("Anonymous unified profiles", rs.get("anonymousUnifiedProfiles", "")),
                ("Data space", rs.get("dataSpaceName", "")),
                ("", ""),
            ]
        )
        label = rs.get("label", "")
        if PROVISIONAL_LABEL.search(label or ""):
            flag(
                "5.1a identity resolution",
                label,
                "Ruleset name looks provisional",
                "Multiple rulesets on one DMO with provisional names usually means only one is real",
                "Confirm which ruleset is authoritative before migrating",
            )
        if rs.get("doesRunAutomatically") is False:
            flag(
                "5.1a identity resolution",
                label,
                "Does not run automatically",
                "A manually run ruleset produces profile counts that drift from reality",
                "Confirm the intended schedule in the target org",
            )
        if rs.get("lastJobStatus") not in (None, "SUCCESS", "Success"):
            flag(
                "5.1a identity resolution",
                label,
                f"Last job status is {rs.get('lastJobStatus')}",
                "Match rate baselines are only valid from a successful run",
                "Re-run before capturing the baseline in 12.1",
            )
    write_rows(t[20], header_rows)
    if len(idres) > 1:
        by_dmo: dict[str, list[str]] = {}
        for rs in idres:
            by_dmo.setdefault(rs.get("objectApiName") or rs.get("secondaryDmo") or "", []).append(
                rs.get("label", "")
            )
        for dmo, labels in by_dmo.items():
            if len(labels) > 1:
                flag(
                    "5.1a identity resolution",
                    dmo,
                    f"{len(labels)} rulesets target the same DMO: {', '.join(labels)}",
                    "Only one ruleset should normally be authoritative per DMO",
                    "Decide which one migrates",
                )

    # ---------- 5.1a match rules (t21)
    match_rows = []
    order = 1
    for rs in idres:
        for mr in rs.get("matchRules") or []:
            crit = mr.get("criteria") or []
            methods = ", ".join(
                sorted({c.get("matchMethodType", "") for c in crit if c.get("matchMethodType")})
            )
            fields = " + ".join(f"{c.get('entityName','')}.{c.get('fieldName','')}" for c in crit)
            flags = " | ".join(
                f"{c.get('fieldName','')}: {c.get('matchMethodType','')}, "
                f"{'case-sensitive' if c.get('caseSensitiveMatch') else 'case-insensitive'}, "
                f"{'matches on blank' if c.get('shouldMatchOnBlank') else 'no blank match'}"
                for c in crit
            )
            match_rows.append(
                (order, f"{rs.get('label','')} :: {mr.get('label','')}", methods, fields, flags, "", "")
            )
            order += 1
    write_rows(t[21], match_rows)
    counts["5.1a identity resolution"] = (
        f"{len(idres)} rulesets, {len(match_rows)} match rules (ssot/identity-resolutions)"
    )

    # ---------- 5.1b reconciliation rules (t22)
    recon_rows = []
    for rs in idres:
        for rec in rs.get("reconciliationRules") or []:
            per_field = []
            for f in rec.get("fields") or []:
                srcs = [s.get("name") for s in (f.get("sources") or []) if s.get("name")]
                if srcs:
                    per_field.append(f"{f.get('fieldName','')}: {' > '.join(srcs)}")
                elif f.get("ruleType"):
                    per_field.append(f"{f.get('fieldName','')}: {f.get('ruleType')}")
            top_srcs = [s.get("name") for s in (rec.get("sources") or []) if s.get("name")]
            if top_srcs:
                per_field.append("ruleset default: " + " > ".join(top_srcs))
            recon_rows.append(
                (
                    f"{rs.get('label','')} :: {rec.get('entityName','')} "
                    f"-> {rec.get('unifiedDmoName','')}",
                    rec.get("ruleType", ""),
                    " | ".join(per_field),
                    "",
                    "",
                )
            )
    write_rows(t[22], recon_rows)
    counts["5.1b reconciliation"] = (
        f"{len(recon_rows)} reconciliation rules (ssot/identity-resolutions)"
    )

    # ---------- 6.1 Calculated insights (t25)
    ci_rows = []
    for name in sorted(ci_items):
        api = ci_items[name]
        m = next((c for c in meta_ci if c.get("name") == name), {})
        sources = sorted(
            {r.get("fromEntity") for r in (m.get("relationships") or []) if r.get("fromEntity")}
        )
        dependents = [p["record"].get("apiName", "") for p in seg_parsed if name in p["objects"]]
        shape = (
            f"{len(m.get('measures') or [])} measures, "
            f"{len(m.get('dimensions') or [])} dimensions, "
            f"{len(sources)} source DMOs"
        )
        ci_rows.append(
            (
                api.get("displayName") or m.get("displayName") or name,
                name,
                api.get("description", ""),
                "; ".join(sources) or NA,
                api.get("publishScheduleInterval", NA),
                NA,
                "; ".join(dependents) or NONE_DETECTED,
                shape,
                "",
                "",
            )
        )
        if api.get("publishScheduleInterval") in (None, "", "NOT_SCHEDULED"):
            flag(
                "6.1 calculated insights",
                name,
                "No refresh schedule",
                "An unscheduled CI serves stale values to anything that reads it",
                "Confirm whether it is still needed",
            )
        last = m.get("latestSuccessfulProcessTime") or ""
        if not last or last.startswith("1970"):
            flag(
                "6.1 calculated insights",
                name,
                "No successful run recorded",
                "A CI that has never run cannot provide a migration baseline",
                "Run it or exclude it",
            )
        if not dependents:
            flag(
                "6.1 calculated insights",
                name,
                "No segment references this CI",
                "CIs with no consumer are often abandoned experiments",
                "Confirm a consumer exists outside segments, otherwise drop",
            )
    write_rows(t[25], ci_rows)
    variants = [c.get("name") for c in meta_ci if c.get("name") not in ci_items]
    counts["6.1 calculated insights"] = (
        f"{len(ci_rows)} CIs (ssot/calculated-insights) with source DMOs, measures and dimensions "
        f"from ssot/metadata?entityType=CalculatedInsight; the Complexity column holds the measure, "
        f"dimension and source counts as evidence for a human rating; the metadata endpoint also "
        f"returns {len(variants)} internal date-rollup entities which are not listed here"
    )

    # ---------- 6.2 Data graphs (t26)
    graph_rows = []
    for g in graphs:
        dev = g.get("developerName", "")
        det = (graph_detail.get(dev) or {}).get("body") or {}
        names: list[str] = []
        edges: list[tuple[str, str]] = []
        walk_graph(g.get("dgObject") or {}, names, edges)
        uniq = list(dict.fromkeys(names))
        refresh = []
        sched = (det.get("fullRefreshConfig") or {}).get("schedule") or {}
        if sched:
            refresh.append(f"full refresh every {sched.get('frequency')} {sched.get('timeGranularity')}")
        inc = det.get("incrementalRefreshConfig") or {}
        if "enabled" in inc:
            refresh.append(f"incremental {'on' if inc.get('enabled') else 'off'}")
        if det.get("isRealTimeToggleEnabled") is not None:
            refresh.append(f"real time {'on' if det.get('isRealTimeToggleEnabled') else 'off'}")
        if det.get("isRecordCachingDisabled") is not None:
            refresh.append(f"record caching {'off' if det.get('isRecordCachingDisabled') else 'on'}")
        zc = [n for n in uniq if n in zero_copy_dmos]
        graph_rows.append(
            (
                det.get("label") or dev,
                dev,
                "; ".join(uniq) or NA,
                "; ".join(f"{a} -> {b}" for a, b in edges) or NA,
                "; ".join(zc) if zc else "N",
                " · ".join(refresh) or NA,
                "",
                "",
                "",
            )
        )
    write_rows(t[26], graph_rows)
    counts["6.2 data graphs"] = (
        f"{len(graph_rows)} data graphs (ssot/data-graphs/metadata for the object tree, "
        f"ssot/data-graphs/<devName> for refresh config); zero-copy is derived by tracing each "
        f"graph DMO back through its DLO mapping to a stream with dataAccessMode DIRECT_ACCESS"
    )

    # ---------- 7.1 Segment inventory (t27)
    dest_by_segment: dict[str, list[str]] = {}
    for a in activations:
        sid = a.get("segmentId")
        tgt = targets.get(a.get("activationTargetId"), {})
        label = a.get("activationTargetName") or tgt.get("name") or ""
        platform = tgt.get("platformName")
        if platform and platform != label:
            label = f"{label} ({platform})"
        if sid:
            dest_by_segment.setdefault(sid, []).append(label)
    seg_rows = []
    for p in seg_parsed:
        s = p["record"]
        sched = s.get("publishScheduleInfo") or {}
        sched_text = (
            " ".join(
                str(x)
                for x in (
                    sched.get("frequency"),
                    f"interval={sched.get('interval')}" if sched.get("interval") else "",
                    "days=" + ",".join(sched.get("daysOfWeek") or []) if sched.get("daysOfWeek") else "",
                )
                if x
            )
            or NA
        )
        seg_rows.append(
            (
                s.get("displayName", ""),
                s.get("apiName", ""),
                "",
                s.get("segmentOnApiName", ""),
                s.get("lastSegmentMemberCount", NA),
                s.get("publishInterval", NA),
                sched_text,
                "",
                "; ".join(dest_by_segment.get(s.get("marketSegmentId"), [])),
                "",
                "",
            )
        )
        nm = s.get("apiName", "")
        if not s.get("lastSegmentMemberCount"):
            flag(
                "7.1 segments",
                nm,
                "Population is zero",
                "An empty segment is usually a test, but may be a broken definition",
                "Confirm before migrating",
            )
        if s.get("publishInterval") in ("NO_REFRESH", None, ""):
            flag(
                "7.1 segments",
                nm,
                "Publish interval is NO_REFRESH",
                "The segment is not being refreshed, so downstream activations serve a frozen audience",
                "Confirm the intended cadence in the target org",
            )
        if PROVISIONAL_LABEL.search(nm):
            flag(
                "7.1 segments",
                nm,
                "Name looks provisional",
                "Provisional names usually mean experiments that should not migrate",
                "Confirm with the segment owner",
            )
        if not dest_by_segment.get(s.get("marketSegmentId")):
            flag(
                "7.1 segments",
                nm,
                "No activation targets this segment",
                "A segment with no destination has no delivery value unless read another way",
                "Confirm whether it is used outside activations",
            )
    write_rows(t[27], seg_rows)
    counts["7.1 segment inventory"] = (
        f"{len(seg_rows)} segments (ssot/segments); destinations joined from ssot/activations on "
        f"marketSegmentId equals segmentId"
    )

    # ---------- 7.2 Segment logic (t28)
    logic_rows = []
    for p in seg_parsed:
        s = p["record"]
        cis = sorted(o for o in p["objects"] if o.endswith("__cio"))
        graph_hits = sorted(
            dev for dev, objs in graph_objects.items() if p["objects"] & objs
        )
        logic_rows.append(
            (
                s.get("apiName", ""),
                describe_criteria(p["include"]) or NA,
                describe_criteria(p["exclude"]),
                "; ".join(cis) or NONE_DETECTED,
                "; ".join(graph_hits) or NONE_DETECTED,
                "",
                "",
            )
        )
    write_rows(t[28], logic_rows)
    counts["7.2 segment logic"] = (
        f"{len(logic_rows)} segments decoded from includeCriteria and excludeCriteria, which "
        f"arrive HTML-escaped; CI and graph dependencies are the objects referenced in the criteria"
    )

    # ---------- 7.3 Segment counts (t29)
    live = sum(1 for s in segments if s.get("segmentStatus") == "ACTIVE")
    put(t[29], 1, 1, live)
    dep_ci = sum(1 for p in seg_parsed if any(o.endswith("__cio") for o in p["objects"]))
    put(t[29], 5, 1, dep_ci)

    # ---------- 8.1a / 8.1b Activations (t32, t33)
    seg_api_by_id = {s.get("marketSegmentId"): s.get("apiName") for s in segments}
    act_rows = []
    pay_rows = []
    for a in sorted(activations, key=lambda x: x.get("name", "")):
        tgt = targets.get(a.get("activationTargetId"), {})
        dest = a.get("activationTargetName", "")
        if tgt.get("platformName"):
            dest = f"{dest} ({tgt['platformName']})"
        act_rows.append(
            (
                a.get("name", ""),
                a.get("developerName", ""),
                dest,
                "",
                "",
                a.get("activationTargetId", ""),
                seg_api_by_id.get(a.get("segmentId")) or a.get("segmentApiName", ""),
                "",
                "",
            )
        )
        det = (act_detail.get(a.get("id")) or act_detail.get(a.get("developerName")) or {}).get(
            "body"
        ) or {}
        attrs = schema_leaves(det.get("activationRecordSchema", ""))
        owner_id = (det.get("createdBy") or {}).get("id")
        owner = users.get(owner_id, owner_id or "")
        pay_rows.append(
            (
                a.get("developerName", ""),
                f"{len(attrs)}: " + "; ".join(attrs) if attrs else NA,
                NA,
                NA,
                a.get("refreshType", NA),
                "",
                f"{owner} (record creator)" if owner else "",
            )
        )
        if a.get("status") not in (None, "ACTIVE"):
            flag(
                "8.1 activations",
                a.get("name", ""),
                f"Status is {a.get('status')}",
                "A failing activation cannot be baselined and may indicate a broken destination credential",
                "Fix or exclude, and confirm the destination account is still valid",
            )
        if PROVISIONAL_LABEL.search(a.get("name", "") or ""):
            flag(
                "8.1 activations",
                a.get("name", ""),
                "Name looks provisional",
                "Provisional names usually mean experiments that should not migrate",
                "Confirm with the activation owner",
            )
    write_rows(t[32], act_rows)
    write_rows(t[33], pay_rows)
    counts["8.1 activations"] = (
        f"{len(act_rows)} activations (ssot/activations) against {len(targets)} activation targets "
        f"(ssot/activation-targets); the attribute list in 8.1b is the payload schema from "
        f"ssot/activations/<id>, and Owner is the record creator resolved through a single SOQL "
        f"query on User, not necessarily the business owner"
    )

    # ---------- 11.3 Search indexes (t42)
    idx_rows = []
    for ix in indexes:
        chunk_cfgs = (ix.get("chunkingConfiguration") or {}).get("fieldLevelConfigurations", []) or []
        fields = sorted(
            {
                d.get("dmoFieldDeveloperName")
                for cfg in chunk_cfgs
                for d in (cfg.get("decorators") or [])
                if d.get("dmoFieldDeveloperName")
            }
        )
        index_type = " · ".join(x for x in (ix.get("searchType"), ix.get("processingType")) if x)
        emb = (ix.get("vectorEmbeddingConfiguration") or {}).get("embeddingModel") or {}
        dims = "; ".join(f"{u.get('id')}={u.get('value')}" for u in (emb.get("userValues") or []))
        model = " · ".join(x for x in (emb.get("id"), dims) if x)
        idx_rows.append(
            (
                ix.get("label") or ix.get("developerName", ""),
                ix.get("developerName", ""),
                index_type or NA,
                ix.get("sourceDmoDeveloperName", NA),
                "; ".join(fields) or NA,
                model or NA,
                "",
                NA,
                "",
                "",
            )
        )
    write_rows(t[42], idx_rows)
    counts["11.3 search indexes"] = (
        f"{len(idx_rows)} search index definitions (ssot/search-index)"
    )

    # ---------- 0.3 Extract provenance (t6)
    api = prov.get("apiVersion", "")
    prov_rows = [
        ("Data space & scope", "ssot/data-spaces, ssot/data-spaces/{name}/members", "data-spaces.json, data-space-members.json"),
        ("Streams & connections", "ssot/data-streams?includeMappings=true, ssot/connections", "data-streams.json, connections.json"),
        ("DLOs", "ssot/data-lake-objects, ssot/metadata?entityType=DataLakeObject", "data-lake-objects.json, metadata-dlo.json"),
        ("DMOs & mappings", "ssot/metadata?entityType=DataModelObject, ssot/data-model-objects, ssot/data-model-object-mappings, Tooling MktDataModelObject", "metadata-dmo.json, data-model-objects-catalogue.json, dmo-mappings.json, tooling-dmo.json"),
        ("IR ruleset", "ssot/identity-resolutions", "identity-resolutions.json"),
        ("CIs & data graphs", "ssot/calculated-insights, ssot/metadata?entityType=CalculatedInsight, ssot/data-graphs/metadata, ssot/data-graphs/{name}", "calculated-insights.json, metadata-ci.json, data-graphs.json, data-graph-details.json"),
        ("Segments", "ssot/segments", "segments.json"),
        ("Activations", "ssot/activations, ssot/activations/{id}, ssot/activation-targets", "activations.json, activation-details.json, activation-targets.json"),
        ("Retrieve / search indexes", "ssot/search-index", "search-index.json"),
    ]
    for i, (_, how, evidence) in enumerate(prov_rows, start=1):
        put(t[6], i, 1, f"Data Cloud REST API {api} GET {how}")
        put(t[6], i, 2, extract_date)
        put(t[6], i, 3, evidence)
    counts["0.3 extract provenance"] = (
        f"9 object classes, each naming the endpoint and the cache file that holds the response"
    )

    # ---------- 1.3 Object reconciliation (t8)
    extracted = {
        1: len(scope_rows),
        2: len(stream_rows),
        3: len(dlo_rows),
        4: len(dmo_rows),
        5: len(idres),
        6: len(ci_rows),
        7: len(graph_rows),
        8: len(seg_rows),
        9: len(act_rows),
        10: "",
        11: len(idx_rows),
    }
    for row, n in extracted.items():
        put(t[8], row, 3, n)
        put(t[8], row, 4, n)
        put(t[8], row, 10, "Y")
    counts["1.3 object reconciliation"] = (
        "Extracted and Documented are the row counts written into each section, so they reconcile "
        "by construction; Migrate, Drop and the packaging columns are decisions and stay blank"
    )
    # 10.1 remains manual because consumer inventory spans systems beyond Data Cloud.

    # ---------- 12.1 Test baselines (t46)
    total_unified = sum(
        rs.get("totalUnifiedProfiles") or 0 for rs in idres if isinstance(rs.get("totalUnifiedProfiles"), (int, float))
    )
    rates = "; ".join(
        f"{rs.get('label','')}: {rs.get('consolidationRate','')}%" for rs in idres
    )
    seg_pops = "; ".join(
        f"{p['record'].get('apiName','')}: {p['record'].get('lastSegmentMemberCount','')}"
        for p in seg_parsed
    )
    filt = (
        "; ".join(f"{a} where {f}" for a, f in filtered_members)
        or "no data space filter applied"
    )
    baselines = {
        1: f"{total_dlo_records:,} rows across {sum(1 for r in dlo_rows if r[4] != NA)} streamed DLOs (per-DLO values in 3.1)",
        2: f"{total_unified:,} unified profiles across {len(idres)} rulesets",
        3: rates or NA,
        4: seg_pops or NA,
        5: NA,
        6: "; ".join(f"{r[0]}: {r[1].split(':')[0]} attributes" for r in pay_rows) or NA,
        7: NA,
        8: filt,
        9: NA,
        10: NA,
    }
    for row, val in baselines.items():
        put(t[46], row, 2, val)
    counts["12.1 test baselines"] = (
        f"Source values are the measurements taken at extract time ({prov.get('fetchedAtUtc','')} "
        f"UTC); tolerance, method and sign-off are for the test lead"
    )

    # ---------- 12.6 Evidence pack (t51)
    cache_files = sorted(p.name for p in CACHE.glob("*.json"))
    put(t[51], 1, 1, f".data-space-analysis-cache/ ({len(cache_files)} JSON responses, one per endpoint)")
    put(t[51], 1, 3, "1")
    put(t[51], 1, 4, extract_date)
    put(t[51], 5, 1, ", ".join(cache_files))
    put(t[51], 5, 3, "1")
    put(t[51], 5, 4, extract_date)
    counts["12.6 evidence pack"] = (
        f"{len(cache_files)} cached endpoint responses listed as the raw metadata extract"
    )

    counts["data transforms (no template section)"] = (
        f"{sum(1 for x in transforms if x.get('type') == 'BATCH')} batch and "
        f"{sum(1 for x in transforms if x.get('type') == 'STREAMING')} streaming transforms exist "
        f"(ssot/data-transforms); the template has no section for them"
    )
    for x in transforms:
        if x.get("lastRunStatus") not in (None, "", "SUCCESS", "Success"):
            flag(
                "data transforms",
                x.get("label") or x.get("name", ""),
                f"Last run status is {x.get('lastRunStatus')}",
                "The template has no section for transforms, so a failing one is easily missed",
                "Decide whether it is in scope and record it manually",
            )

    return counts, observations


# ---------------------------------------------------------------- appendices


def collapse(observations: list[tuple], threshold: int = 3, sample: int = 12) -> list[tuple]:
    """Group identical observations so the appendix reads as findings, not a log.

    One row per distinct (area, observation, reason, action). Objects are listed while
    the list stays readable, then summarised by count.
    """
    groups: dict[tuple, list[str]] = {}
    for area, obj, observation, why, action in observations:
        groups.setdefault((area, observation, why, action), []).append(obj)

    rows = []
    for (area, observation, why, action), objs in groups.items():
        if len(objs) <= threshold:
            for obj in objs:
                rows.append((area, obj, observation, why, action))
        else:
            shown = "; ".join(sorted(objs)[:sample])
            more = len(objs) - min(sample, len(objs))
            obj_cell = f"{len(objs)} objects: {shown}" + (f" ... and {more} more" if more else "")
            rows.append((area, obj_cell, observation, why, action))
    rows.sort(key=lambda r: (r[0], r[2]))
    return rows


def appendices(doc, prov: dict, counts: dict, observations: list[tuple]) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A — Fill provenance", level=1)
    doc.add_paragraph(
        f"Org {prov.get('orgId','')} · instance {prov.get('instanceUrl','')} · "
        f"API {prov.get('apiVersion','')} · fetched {prov.get('fetchedAtUtc','')} UTC · "
        f"{prov.get('apiCalls','?')} API calls in {prov.get('elapsedSeconds','?')} s"
    )
    doc.add_paragraph(
        "Every populated cell comes from the endpoint named below, read with the Salesforce CLI "
        "session of the user who ran the extract. Blank cells are migration decisions or human "
        "inputs and are listed in Appendix B. Cells reading NOT AVAILABLE FROM API have no source "
        "endpoint in this API version."
    )
    for k in sorted(counts):
        doc.add_paragraph(f"{k}: {counts[k]}", style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Appendix B — Human input register", level=1)
    doc.add_paragraph(
        "Every column the extract deliberately leaves alone, why, and where the answer comes from. "
        "Rows marked yes in the last column change migration scope, effort or legal exposure, so "
        "they should be filled by one person and reviewed by a second."
    )
    add_table(
        doc,
        ["§", "Column / content", "Why it is not automated", "Where the answer comes from", "Needs 4 eyes"],
        HUMAN_INPUTS,
    )

    doc.add_page_break()
    doc.add_heading("Appendix C — Observations needing review", level=1)
    grouped = collapse(observations)
    by_area: dict[str, int] = {}
    for area, *_ in observations:
        by_area[area] = by_area.get(area, 0) + 1
    doc.add_paragraph(
        f"{len(observations)} observations across {len(by_area)} areas, grouped into "
        f"{len(grouped)} findings. Each one is a fact read from the org that a human should "
        "interpret before the record is signed off. Nothing here is a conclusion: no object has "
        "been marked in or out of scope, and no value in the document has been changed on the "
        "strength of these notes."
    )
    doc.add_paragraph("By area: " + "; ".join(f"{k} ({v})" for k, v in sorted(by_area.items())))
    add_table(
        doc,
        ["Area", "Object(s)", "Observation", "Why it needs a human", "Suggested action"],
        grouped,
    )


def main() -> None:
    import argparse
    import sys

    sys.stdout.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(
        description="Generate an analysis document for one specific data space."
    )
    parser.add_argument(
        "--space",
        help="Data space name (API name), for example 'default'",
    )
    parser.add_argument(
        "--all-spaces",
        action="store_true",
        help="Generate one document per data space found in cache",
    )
    parser.add_argument(
        "--output-dir",
        default=str(TEMPLATE.parent),
        help="Directory for generated docx files",
    )
    args = parser.parse_args()
    if not args.all_spaces and not args.space:
        parser.error("Provide --space <name> or --all-spaces")

    prov = load("_provenance")
    org = prov.get("orgId", "unknown")
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    spaces = [s.get("name") for s in records("data-spaces") if s.get("name")]
    targets = spaces if args.all_spaces else [args.space]
    for target in targets:
        safe_space = re.sub(r"[^A-Za-z0-9_.-]+", "_", target)
        out = out_dir / (
            f"DataCloud_DataSpace_Analysis_Record_{safe_space}_LIVE_{org}_"
            f"{date.today().strftime('%Y%m%d')}.docx"
        )
        doc = docx.Document(str(TEMPLATE))
        counts, observations = fill(doc, prov, target)
        appendices(doc, prov, counts, observations)
        doc.save(str(out))
        print(f"org={org} space={target} fetched={prov.get('fetchedAtUtc')}")
        print(f"  observations flagged for review: {len(observations)}")
        print(f"Saved: {out}")


if __name__ == "__main__":
    main()
